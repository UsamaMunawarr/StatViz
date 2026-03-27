import streamlit as st
import pandas as pd
import numpy as np
from fpdf import FPDF
import io
import matplotlib.pyplot as plt
import seaborn as sns
import base64
import tempfile
import os
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

class PDFReport(FPDF):
    def header(self):
        self.set_font("Times", "B", 18)
        self.set_text_color(0, 102, 204) # Deep Blue
        self.cell(0, 10, "StatViz - Premium Analytical Report", align="C", border=0, fill=0)
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font("Times", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

def run():
    st.header("🖨️ Automated Document Generation")
    
    if st.session_state.data is not None:
        df = st.session_state.data
        
        st.markdown("Generate a comprehensive, downloadable **PDF or MS Word Document** summarizing your dataset. This intelligent engine analyzes your active data and compiles descriptive stats, correlation matrices, missing value diagnostics, and variable distributions into a highly professional format.")
        
        report_title = st.text_input("Enter Document Title", value="Exploratory Data Analysis & Statistical Summary")
        include_plots = st.checkbox("Include Visualizations (Histograms, Heatmaps & Boxplots)", value=True)
        
        st.divider()
        col1, col2 = st.columns(2)
        
        # ------------------------ PDF GENERATOR ------------------------
        with col1:
            st.write("### 📄 PDF Document")
            st.write("Export a beautifully formatted, read-only PDF file in Times New Roman.")
            if st.button("🚀 Generate PDF Report", use_container_width=True):
                with st.spinner("Compiling Statistical PDF Report..."):
                    try:
                        pdf = PDFReport()
                        pdf.add_page()
                        
                        # Title
                        pdf.set_font("Times", "B", 22)
                        pdf.set_text_color(10, 22, 40)
                        pdf.cell(0, 10, report_title, align="C")
                        pdf.ln(15)
                        
                        # 1. Dataset Overview
                        pdf.set_font("Times", "B", 14)
                        pdf.set_text_color(0, 0, 0)
                        pdf.cell(0, 10, "1. Executive Dataset Overview")
                        pdf.ln(8)
                        
                        pdf.set_font("Times", "", 12)
                        pdf.multi_cell(0, 6, "This section summarizes the high-level architecture of the dataset under investigation. It outlines the core dimensions, structural integrity (missing values), and computational footprint of the active data environment.")
                        pdf.ln(4)
                        
                        pdf.set_font("Times", "B", 11)
                        pdf.cell(0, 8, f"Total Observations (Rows): {df.shape[0]}")
                        pdf.ln(6)
                        pdf.cell(0, 8, f"Total Features (Columns): {df.shape[1]}")
                        pdf.ln(6)
                        pdf.cell(0, 8, f"Total Missing Data Cells: {df.isnull().sum().sum()}")
                        pdf.ln(6)
                        pdf.cell(0, 8, f"Memory Usage: {df.memory_usage(deep=True).sum() / 1024**2:.2f} MB")
                        pdf.ln(12)
                        
                        # 2. Missing Values Analysis
                        missing_stats = df.isnull().sum()
                        missing_stats = missing_stats[missing_stats > 0]
                        
                        pdf.set_font("Times", "B", 14)
                        pdf.cell(0, 10, "2. Data Quality & Missingness Diagnosis")
                        pdf.ln(8)
                        pdf.set_font("Times", "", 12)
                        if missing_stats.empty:
                            pdf.cell(0, 8, "Excellent data integrity: No missing values were detected in the active dataset.")
                            pdf.ln(10)
                        else:
                            pdf.multi_cell(0, 6, "The following variables contain missing records. Understanding missingness is crucial for statistical validity and avoiding bias during Machine Learning modeling.")
                            pdf.ln(4)
                            for col, val in missing_stats.items():
                                perc = (val / len(df)) * 100
                                pdf.cell(0, 6, f"- {col}: {val} missing ({perc:.2f}%)")
                                pdf.ln(6)
                            pdf.ln(4)
                            
                        # 3. Descriptive Stats for Numerics
                        num_cols = df.select_dtypes(include=[np.number]).columns
                        if len(num_cols) > 0:
                            pdf.set_font("Times", "B", 14)
                            pdf.cell(0, 10, "3. Core Descriptive Statistics")
                            pdf.ln(8)
                            pdf.set_font("Times", "", 12)
                            pdf.multi_cell(0, 6, "This table provides the central tendencies, dispersion, and distributional shape of the continuous features. It acts as the mathematical baseline for further parametric testing.")
                            pdf.ln(4)
                            
                            desc = df[num_cols].describe().round(2)
                            
                            pdf.set_font("Times", "B", 9)
                            # Header Row
                            col_width = 25
                            pdf.cell(col_width, 8, "Statistic", border=1)
                            plot_cols = num_cols[:6] # limit to 6 cols to fit on page width
                            for col in plot_cols:
                                pdf.cell(col_width, 8, str(col)[:10], border=1)
                            pdf.ln(8)
                            
                            pdf.set_font("Times", "", 9)
                            for stat_idx in desc.index:
                                pdf.cell(col_width, 8, str(stat_idx).capitalize(), border=1)
                                for col in plot_cols:
                                    pdf.cell(col_width, 8, str(desc.loc[stat_idx, col]), border=1)
                                pdf.ln(8)
                            pdf.ln(15)
                            
                            # Add Distribution Plots
                            if include_plots:
                                pdf.add_page()
                                pdf.set_font("Times", "B", 14)
                                pdf.cell(0, 10, "4. Advanced Visual Analytics")
                                pdf.ln(8)
                                pdf.set_font("Times", "", 12)
                                pdf.multi_cell(0, 6, "Statistical visualizations detailing the empirical distribution of continuous variables and their inter-item correlations. Used for identifying skewness, multimodality, and multicollinearity.")
                                pdf.ln(8)
                                
                                # Correlation Matrix
                                if len(num_cols) >= 2:
                                    plt.figure(figsize=(7, 5))
                                    corr = df[num_cols].corr()
                                    sns.heatmap(corr, annot=True, cmap="coolwarm", fmt=".2f", linewidths=0.5)
                                    plt.title("Pearson Correlation Matrix")
                                    plt.tight_layout()
                                    temp_corr = os.path.join(tempfile.gettempdir(), "corr_matrix.png")
                                    plt.savefig(temp_corr, dpi=120)
                                    plt.close()
                                    
                                    pdf.image(temp_corr, w=150)
                                    pdf.ln(10)
                                    os.remove(temp_corr)
                                
                                # Distributions
                                for col in plot_cols:
                                    plt.figure(figsize=(6, 4))
                                    sns.histplot(df[col].dropna(), kde=True, color="#00d4aa")
                                    plt.title(f"Density Distribution of {col}")
                                    plt.tight_layout()
                                    
                                    temp_path = os.path.join(tempfile.gettempdir(), f"{col}_dist.png")
                                    plt.savefig(temp_path, dpi=120)
                                    plt.close()
                                    
                                    if pdf.get_y() > 200:
                                        pdf.add_page()
                                    pdf.image(temp_path, w=120)
                                    pdf.ln(5)
                                    os.remove(temp_path)
                                
                        pdf_output = bytes(pdf.output())
                        
                        st.success("✅ PDF Generation Complete!")
                        st.download_button(
                            label="📥 Download PDF Document",
                            data=pdf_output,
                            file_name="StatViz_Pro_Report.pdf",
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"Failed to generate PDF Report. Error: {str(e)}")
        
        # ------------------------ WORD GENERATOR ------------------------
        with col2:
            st.write("### 📝 MS Word Document")
            st.write("Export an editable `.docx` file in Times New Roman. Fully customizable.")
            if st.button("🚀 Generate Word Report", use_container_width=True):
                with st.spinner("Compiling Statistical Word Document..."):
                    try:
                        doc = Document()
                        
                        # Set default font to Times New Roman
                        style = doc.styles['Normal']
                        style.font.name = 'Times New Roman'
                        style.font.size = Pt(12)
                        
                        # Title
                        title = doc.add_heading(report_title, 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 1. Dataset Overview
                        doc.add_heading('1. Executive Dataset Overview', level=1)
                        doc.add_paragraph("This section summarizes the high-level architecture of the dataset under investigation. It outlines the core dimensions, structural integrity (missing values), and computational footprint of the active data environment.")
                        
                        doc.add_paragraph(f"Total Observations (Rows): {df.shape[0]}", style='List Bullet')
                        doc.add_paragraph(f"Total Features (Columns): {df.shape[1]}", style='List Bullet')
                        doc.add_paragraph(f"Total Missing Data Cells: {df.isnull().sum().sum()}", style='List Bullet')
                        doc.add_paragraph(f"Memory Usage: {df.memory_usage(deep=True).sum() / 1024**2:.2f} MB", style='List Bullet')
                        
                        # 2. Missing Values Analysis
                        missing_stats = df.isnull().sum()
                        missing_stats = missing_stats[missing_stats > 0]
                        doc.add_heading('2. Data Quality & Missingness Diagnosis', level=1)
                        if missing_stats.empty:
                            doc.add_paragraph("Excellent data integrity: No missing values were detected in the active dataset.")
                        else:
                            doc.add_paragraph("The following variables contain missing records. Understanding missingness is crucial for statistical validity and avoiding bias during Machine Learning modeling.")
                            for col, val in missing_stats.items():
                                perc = (val / len(df)) * 100
                                doc.add_paragraph(f"{col}: {val} missing ({perc:.2f}%)", style='List Bullet')
                                
                        # 3. Descriptive Stats
                        num_cols = df.select_dtypes(include=[np.number]).columns
                        if len(num_cols) > 0:
                            doc.add_heading('3. Core Descriptive Statistics', level=1)
                            doc.add_paragraph("This table provides the central tendencies, dispersion, and distributional shape of the continuous features. It acts as the mathematical baseline for further parametric testing.")
                            
                            desc = df[num_cols].describe().round(2)
                            plot_cols = num_cols[:6]
                            
                            table = doc.add_table(rows=1, cols=len(plot_cols)+1)
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Statistic'
                            for i, col in enumerate(plot_cols):
                                hdr_cells[i+1].text = str(col)[:10]
                                
                            for stat_idx in desc.index:
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(stat_idx).capitalize()
                                for i, col in enumerate(plot_cols):
                                    row_cells[i+1].text = str(desc.loc[stat_idx, col])
                                    
                            # Visualizations
                            if include_plots:
                                doc.add_page_break()
                                doc.add_heading('4. Advanced Visual Analytics', level=1)
                                doc.add_paragraph("Statistical visualizations detailing the empirical distribution of continuous variables and their inter-item correlations. Used for identifying skewness, multimodality, and multicollinearity.")
                                
                                if len(num_cols) >= 2:
                                    plt.figure(figsize=(7, 5))
                                    corr = df[num_cols].corr()
                                    sns.heatmap(corr, annot=True, cmap="coolwarm", fmt=".2f", linewidths=0.5)
                                    plt.title("Pearson Correlation Matrix")
                                    plt.tight_layout()
                                    temp_corr = os.path.join(tempfile.gettempdir(), "word_corr.png")
                                    plt.savefig(temp_corr, dpi=120)
                                    plt.close()
                                    
                                    doc.add_picture(temp_corr, width=Inches(5.0))
                                    os.remove(temp_corr)
                                
                                for col in plot_cols:
                                    plt.figure(figsize=(6, 4))
                                    sns.histplot(df[col].dropna(), kde=True, color="#00d4aa")
                                    plt.title(f"Density Distribution of {col}")
                                    plt.tight_layout()
                                    
                                    temp_path = os.path.join(tempfile.gettempdir(), f"word_{col}.png")
                                    plt.savefig(temp_path, dpi=120)
                                    plt.close()
                                    
                                    doc.add_picture(temp_path, width=Inches(4.5))
                                    os.remove(temp_path)
                                    
                        # Save doc to BytesIO
                        word_io = io.BytesIO()
                        doc.save(word_io)
                        word_io.seek(0)
                        
                        st.success("✅ MS Word Generation Complete!")
                        st.download_button(
                            label="📥 Download Word Document (.docx)",
                            data=word_io,
                            file_name="StatViz_Pro_Report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Failed to generate Word document. Error: {str(e)}")

    else:
        st.warning("Please upload data first in the Data Import module.")
